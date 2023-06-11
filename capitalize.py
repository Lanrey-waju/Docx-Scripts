import sys
from docx import Document


def capitalize_word(file_path, word):
    # Load document
    doc = Document(file_path)

    # Iterate over the paragraphs
    for paragraph in doc.paragraphs:
        runs = paragraph.runs

        # Iterate over each run
        for i, run in enumerate(runs):
            if word in run.text:
                # Separate the runs with the given word or sequence
                parts = run.text.split(word)
                run.text = parts[0]

                # Capitalize the word sequence
                new_run = paragraph.add_run(word.capitalize())

                paragraph.runs.insert(i + 1, new_run)

                # Maintain the formatting of the other parts
                if len(parts) > 1:
                    new_run = paragraph.add_run(parts[1])
                    new_run.bold = run.bold
                    new_run.italic = run.italic
                    new_run.underline = run.underline
                    paragraph.runs.insert(i + 2, new_run)

    doc.save("capitalized_doc.docx")


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python capitalize.py <file_path> <word>")
        sys.exit(1)

    file_path = sys.argv[1]
    word = sys.argv[2]
    capitalize_word(file_path, word)
