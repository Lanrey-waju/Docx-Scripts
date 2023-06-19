import sys
from docx import Document


def edit_word_in_docx(file_path, word, flag):
    # Load the Word document
    doc = Document(file_path)

    # Iterate over paragraphs in the document
    for paragraph in doc.paragraphs:
        runs = paragraph.runs
        for i, run in enumerate(runs):
            if word in run.text:
                # Split the run's text based on the word
                parts = run.text.split(word)
                run.text = parts[0]

                # Create a new run for the word and edit it
                if flag == "-i":
                    new_run = paragraph.add_run(word)
                    new_run.italic = True
                elif flag == "-c":
                    new_run = paragraph.add_run(word.capitalize())

                # Insert the new run after the original run
                paragraph.runs.insert(i + 1, new_run)

                # If there's text after the word, create another run
                if len(parts) > 1:
                    new_run = paragraph.add_run(parts[1])
                    # Preserve the formatting of the original run
                    new_run.bold = run.bold
                    new_run.italic = run.italic
                    new_run.underline = run.underline
                    paragraph.runs.insert(i + 2, new_run)

    # Save the modified document
    if flag == "-i":
        doc.save("italicized_document.docx")
    elif flag == "-c":
        doc.save("capitalized_document.docx")


if __name__ == "__main__":
    if len(sys.argv) != 4:
        print("Usage: python edit.py <file_path> <word> <flag>")
        sys.exit(1)

    file_path = sys.argv[1]
    word = sys.argv[2]
    flag = sys.argv[3]
    edit_word_in_docx(file_path, word, flag)
